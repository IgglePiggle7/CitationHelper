import gradio as gr
import docx
import re
import os
import requests
import io
from openai import OpenAI

# 获取环境变量 (API Keys)
SERPAPI_KEY = os.environ.get("SERPAPI_KEY", "")
LLM_API_KEY = os.environ.get("LLM_API_KEY", "")
# 默认使用通义千问的 Base URL，如果用其他模型请在此修改或通过环境变量配置
LLM_BASE_URL = os.environ.get("LLM_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1") 

# 功能一：文献引用生成器 (SerpApi + LLM)
def generate_citation(paper_title, selected_styles):
    if not paper_title.strip():
        return "❌ 请输入文献标题。"
    if not selected_styles:
        return "❌ 请至少选择一种引用格式。"
    if not SERPAPI_KEY or not LLM_API_KEY:
        return "❌ 系统未配置 API Key，请在 Hugging Face Secrets 中配置 SERPAPI_KEY 和 LLM_API_KEY。"

    try:
        # 1. 第一步：在 Google Scholar 搜索文献，获取 result_id
        search_url = "https://serpapi.com/search"
        search_params = {
            "engine": "google_scholar",
            "q": paper_title,
            "api_key": SERPAPI_KEY,
            "num": 1
        }
        search_res = requests.get(search_url, params=search_params).json()
        
        organic_results = search_res.get("organic_results", [])
        if not organic_results:
            return "⚠️ 未在 Google Scholar 找到该文献，请检查标题是否拼写正确。"
            
        result_id = organic_results[0].get("result_id")
        real_title = organic_results[0].get("title", paper_title)
        
        # 2. 第二步：获取该文献的多种引用格式原始数据
        cite_params = {
            "engine": "google_scholar_cite",
            "q": result_id,
            "api_key": SERPAPI_KEY
        }
        cite_res = requests.get(search_url, params=cite_params).json()
        raw_citations = cite_res.get("citations", [])
        links = cite_res.get("links", [])
        
        # 将原始数据转为字符串，喂给大模型
        raw_data_str = f"文献标题: {real_title}\n"
        raw_data_str += "SerpApi 提供的基础格式:\n" + str(raw_citations) + "\n"
        raw_data_str += "SerpApi 提供的链接(含BibTeX等):\n" + str(links)

        # 3. 第三步：调用大模型进行精准格式化
        client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)
        
        styles_str = ", ".join(selected_styles)
        prompt = f"""
你是一个严格的学术参考文献格式生成专家。
任务：基于以下 Google Scholar 提供的原始引用数据，**仅**生成用户指定的格式，并直接以标准的 **Markdown** 格式输出。

【用户指定的格式】: {styles_str}

【原始引用数据】:
{raw_data_str}

【格式规范库】:
1. **GB/T 7714-2015**: 
   - 格式：AUTHOR A. Title[J]. Journal, Year, Vol(Issue): Pages.
   - 必须提取卷号和期号，如 27(11): 22-30。
   - 除了结尾的"."外，"."、","、":"等英文的标点符号后面都加空格，全部都是半角输入。
   - 作者规则: 3位以内全部列出；超过3位，仅列出前3位，后加 ", et al." (英文) 或 ", 等." (中文)，绝对不能截断中间的作者。
   - 文献类型标识：普通图书 [M] 期刊文章 [J] 论文集 [C] 学位论文 [D] 报告 [R] 标准 [S] 电子资源 [EB/OL]
   - 作者署名：中文作者：姓氏全拼+名字首字母（如：张华 → Zhang H），3人以内全部列出，超过3人则前3人+等或et al.。外文作者：姓氏全大写+名字首字母（如：SMITH J），3人以内全部列出，超过3人则前3人+等或et al.。
   - 严禁在结尾添加 "URL: https://..."，除非该文章是纯网络首发且没有页码。如果有 DOI，仅在 APA/MLA 中展示，国标通常省略 DOI 除非特定要求。
2. **APA7**:
   - 格式：Author, A. (Year). Title. *Journal, Vol*(Issue), Pages. (期刊名和卷号用 Markdown 斜体)
   - **作者规则**: **列出多达 20 位作者**。只有人数超过20个的时候才能在最后写 "et al." ，但绝对不能截断中间的作者。
   - **作者格式**: Author, A., & Author, B. (Year). ... (注意最后一位作者前用 "&")。
3. **MLA9**:
   - 格式：Author, A. "Title." *Journal*, vol. V, no. I, Year, pp. P. (期刊名斜体)
   - **作者规则**: **3位或以上作者，仅列出第一位作者，后加 ", et al."**。
   - **作者格式**: Author, A., et al. "Title." ...
   - 2位作者则写: Author, A., and B. Author.
4. **IEEE**:
   - 格式：[1] J. K. Author, "Title of paper," *Abbrev. Title of Journal*, vol. x, no. x, pp. xxx-xxx, Abbrev. Month, Year.
   - 注意：作者名字首字母在前；文章名加双引号；期刊名斜体。
   - **作者规则**: 列出所有作者。如果作者人数极多(>6)，可列第一位 + "et al."，绝对不能截断中间的作者。
   - 作者格式: [1] J. K. Author, "Title," ... (名缩写在前，姓在后)。
5. **Chicago (Bibliography)**:
   - 格式：Author, First. "Title of Article." *Journal Title* Volume, no. Issue (Year): Pages.
   - **作者规则**: 参考书目(Bibliography)中，列出多达 10 位作者；超过 10 位列出前 7 位 + "et al."。
   
【特殊数据处理】:
如果输入元数据中的作者列表已经被截断（例如包含 "..." 或 "Horizontal ellipsis"），且无法获取完整名单：
- GB/T 和 MLA: 请确保符合上述 "et al." 规则。
- APA 和 IEEE: 列出所有已知作者，并在末尾保留 "et al." 以示严谨，不要生造名字。

【输出要求】:
1. 严禁输出用户未选择的格式。
2. 如果原始数据缺失卷期号，尽力从 BibTeX 链接或其他格式中推断，若无则留空。
3. 请直接输出结果，不要包含任何如“好的”、“为您生成”等闲聊废话。使用 Markdown 的 `### 格式名` 作为小标题。
"""
        # 注意这里模型名：如果用阿里填 qwen-plus/qwen-max，如果用其他平台请换成对应模型名
        response = client.chat.completions.create(
            model="qwen-plus", 
            messages=[
                {"role": "system", "content": "You are a professional academic citation assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        return response.choices[0].message.content

    except Exception as e:
        return f"❌ 运行过程中发生错误: {str(e)}"

# 功能二：Word 文献序号重排 (纯 Python 处理)
def process_reorder(doc_file_path, ref_text):
    if not doc_file_path or not ref_text.strip():
        return None, "❌ 错误：请先上传 Word 文档并粘贴参考文献！"
    try:
        doc = docx.Document(doc_file_path)
        full_text = [para.text for para in doc.paragraphs]
        doc_text_str = "\n".join(full_text)

        def expand_range(range_str):
            try:
                range_str = range_str.replace('–', '-')
                if '-' in range_str:
                    s, e = range_str.split('-')
                    return [str(i) for i in range(int(s), int(e) + 1)]
                return [range_str]
            except: return []

        def compress_range(nums):
            if not nums: return ""
            nums = sorted(list(set(nums)))
            ranges, start, prev = [], nums[0], nums[0]
            for x in nums[1:]:
                if x == prev + 1: prev = x
                else:
                    ranges.append(str(start) if start == prev else f"{start}-{prev}")
                    start = prev = x
            ranges.append(str(start) if start == prev else f"{start}-{prev}")
            return ", ".join(ranges)

        matches = re.findall(r'\[([\d\s,，\-–]+)\]', doc_text_str)
        doc_order = []
        seen = set()
        for content in matches:
            for part in re.split(r'[,，]', content.replace(" ", "")):
                if not part: continue
                for pid in expand_range(part):
                    if pid.isdigit() and pid not in seen:
                        seen.add(pid); doc_order.append(pid)

        ref_map = {}
        pattern = re.compile(r'^\s*(?:\[(\d+)\]|(\d+)\.)\s*(.*)')
        for line in ref_text.strip().split('\n'):
            match = pattern.match(line.strip())
            if match:
                oid = match.group(1) or match.group(2)
                ref_map[oid] = match.group(3)

        mapping, new_ref_list, current_idx = {}, [], 1
        for old_id in doc_order:
            if old_id in ref_map:
                mapping[old_id] = current_idx
                new_ref_list.append(f"[{current_idx}] {ref_map[old_id]}")
                del ref_map[old_id]; current_idx += 1
                
        for old_id, content in ref_map.items():
            mapping[old_id] = current_idx 
            new_ref_list.append(f"[{current_idx}] {content}")
            current_idx += 1

        def replace_callback(match):
            parts = re.split(r'[,，]', match.group(1))
            new_ids = []
            for part in parts:
                for old_id in expand_range(part.strip()):
                    if old_id in mapping: new_ids.append(mapping[old_id])
                    elif old_id.isdigit(): new_ids.append(int(old_id))
            return f"[{compress_range(new_ids)}]"

        for para in doc.paragraphs:
            if "[" in para.text:
                try:
                    new_text = re.sub(r'\[([\d\s,，\-–]+)\]', replace_callback, para.text)
                    if new_text != para.text: para.text = new_text
                except: pass

        doc.add_paragraph("\n")
        doc.add_paragraph("参考文献 (Auto Generated)").style = 'Heading 1'
        for ref in new_ref_list: doc.add_paragraph(ref)

        output_filename = "Fixed_Document.docx"
        doc.save(output_filename)
        return output_filename, "\n".join(new_ref_list)
        
    except Exception as e:
        return None, f"❌ 处理出错: {str(e)}"

# 网页界面 (Gradio UI)
# 设置主题，让界面看起来更现代、专业
theme = gr.themes.Soft(
    primary_hue="blue",
    neutral_hue="slate",
    font=[gr.themes.GoogleFont("Inter"), "ui-sans-serif", "system-ui", "sans-serif"]
)

with gr.Blocks(theme=theme, title="一站式学术写作神器") as demo:
    gr.Markdown("<center><h1>🎓 一站式学术写作神器</h1></center>")
    gr.Markdown("<center>文献格式一键生成 | Word 引用序号自动重排</center>")
    
    with gr.Tabs():
        # -----------------------------
        # 标签页 1：文献格式生成器
        # -----------------------------
        with gr.TabItem("🔍 1. 引用格式生成器"):
            gr.Markdown("输入论文标题，自动从 Google Scholar 抓取数据并格式化为您需要的标准引用。")
            with gr.Row():
                with gr.Column(scale=1):
                    paper_input = gr.Textbox(label="输入文献标题", placeholder="例如：Attention is all you need")
                    style_choices = gr.CheckboxGroup(
                        choices=["GB/T 7714-2015", "APA7", "MLA9", "IEEE", "Chicago", "BibTeX"],
                        value=["GB/T 7714-2015", "APA7"], # 默认选中两项
                        label="选择目标格式 (可多选)"
                    )
                    gen_btn = gr.Button("🚀 联网生成引用", variant="primary")
                    
                with gr.Column(scale=1):
                    gen_output = gr.Markdown(label="生成结果")
            
            # 绑定生成按钮事件
            gen_btn.click(fn=generate_citation, inputs=[paper_input, style_choices], outputs=gen_output)

        # -----------------------------
        # 标签页 2：Word 乱序重排
        # -----------------------------
        with gr.TabItem("📝 2. Word 引用乱序重排"):
            gr.Markdown("将带有乱序标记的 Word 文档（如 `[5, 1-3]`）和原始参考文献列表上传，工具会自动根据正文顺序重新编号。")
            with gr.Row():
                with gr.Column():
                    file_in = gr.File(label="上传 Word 文档 (.docx)", file_types=[".docx"])
                    text_in = gr.Textbox(label="粘贴乱序的参考文献列表", lines=10, placeholder="[1] 原第一篇...\n[2] 原第二篇...")
                    reorder_btn = gr.Button("⚙️ 开始重排并下载文档", variant="primary")
                    
                with gr.Column():
                    file_out = gr.File(label="📥 下载修改好的 Word 文档")
                    text_out = gr.Textbox(label="📋 预览新参考文献列表", lines=14, interactive=False)

            # 绑定重排按钮事件
            reorder_btn.click(fn=process_reorder, inputs=[file_in, text_in], outputs=[file_out, text_out])

# 启动应用
if __name__ == "__main__":
    demo.launch()