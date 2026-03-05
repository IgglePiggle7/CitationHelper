import docx
import re

def process_reorder(doc_file_path, ref_text):
    """
    处理 Word 文档内的引用标号乱序重排，并输出新的文档。
    """
    if not doc_file_path or not ref_text.strip():
        return None, "❌ 错误：请先上传 Word 文档并粘贴参考文献！"
    try:
        doc = docx.Document(doc_file_path)
        full_text =[para.text for para in doc.paragraphs]
        doc_text_str = "\n".join(full_text)

        def expand_range(range_str):
            try:
                range_str = range_str.replace('–', '-')
                if '-' in range_str:
                    s, e = range_str.split('-')
                    return[str(i) for i in range(int(s), int(e) + 1)]
                return [range_str]
            except: return[]

        def compress_range(nums):
            if not nums: return ""
            nums = sorted(list(set(nums)))
            ranges, start, prev = [], nums[0], nums[0]
            for x in nums[1:]:
                if x == prev + 1: prev = x
                else:
                    ranges.append(str(start) if start == prev else f"{start}-{prev}")
                    start = prev = x
            ranges.append(str(start) if start == prev else f"{start}-{prev}")
            return ", ".join(ranges)

        matches = re.findall(r'\[([\d\s,，\-–]+)\]', doc_text_str)
        doc_order =[]
        seen = set()
        for content in matches:
            for part in re.split(r'[,，]', content.replace(" ", "")):
                if not part: continue
                for pid in expand_range(part):
                    if pid.isdigit() and pid not in seen:
                        seen.add(pid); doc_order.append(pid)

        ref_map = {}
        pattern = re.compile(r'^\s*(?:\[(\d+)\]|(\d+)\.)\s*(.*)')
        for line in ref_text.strip().split('\n'):
            match = pattern.match(line.strip())
            if match:
                oid = match.group(1) or match.group(2)
                ref_map[oid] = match.group(3)

        mapping, new_ref_list, current_idx = {},[], 1
        for old_id in doc_order:
            if old_id in ref_map:
                mapping[old_id] = current_idx
                new_ref_list.append(f"[{current_idx}] {ref_map[old_id]}")
                del ref_map[old_id]; current_idx += 1
                
        for old_id, content in ref_map.items():
            mapping[old_id] = current_idx 
            new_ref_list.append(f"[{current_idx}] {content}")
            current_idx += 1

        def replace_callback(match):
            parts = re.split(r'[,，]', match.group(1))
            new_ids =[]
            for part in parts:
                for old_id in expand_range(part.strip()):
                    if old_id in mapping: new_ids.append(mapping[old_id])
                    elif old_id.isdigit(): new_ids.append(int(old_id))
            return f"[{compress_range(new_ids)}]"

        for para in doc.paragraphs:
            if "[" in para.text:
                try:
                    new_text = re.sub(r'\[([\d\s,，\-–]+)\]', replace_callback, para.text)
                    if new_text != para.text: para.text = new_text
                except: pass

        doc.add_paragraph("\n")
        doc.add_paragraph("参考文献 (Auto Generated)").style = 'Heading 1'
        for ref in new_ref_list: doc.add_paragraph(ref)

        output_filename = "Fixed_Document.docx"
        doc.save(output_filename)
        return output_filename, "\n".join(new_ref_list)
        
    except Exception as e:
        return None, f"❌ 处理出错: {str(e)}"